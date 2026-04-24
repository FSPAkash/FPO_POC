"""Build 3-page FPO Integrated OS sales + technical brief."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ACCENT = RGBColor(0x0F, 0x52, 0x3A)
MUTED = RGBColor(0x4A, 0x4A, 0x4A)
DARK = RGBColor(0x11, 0x11, 0x11)


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, size=22, color=ACCENT, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Calibri"
    return p


def add_subheading(doc, text, size=12, color=ACCENT):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Calibri"
    return p


def add_para(doc, text, size=10, color=DARK, bold=False, italic=False, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Calibri"
    r.bold = bold
    r.italic = italic
    return p


def add_bullet(doc, text, size=10, color=DARK, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    if bold_lead:
        r1 = p.add_run(bold_lead)
        r1.bold = True
        r1.font.size = Pt(size)
        r1.font.color.rgb = color
        r1.font.name = "Calibri"
        r2 = p.add_run(text)
        r2.font.size = Pt(size)
        r2.font.color.rgb = color
        r2.font.name = "Calibri"
    else:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.color.rgb = color
        r.font.name = "Calibri"
    return p


def style_table(table, header_fill="0F523A"):
    table.autofit = True
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                for r in p.runs:
                    r.font.name = "Calibri"
                    r.font.size = Pt(9)
                    if i == 0:
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        r.bold = True
                    else:
                        r.font.color.rgb = DARK
            if i == 0:
                set_cell_shading(cell, header_fill)
            elif i % 2 == 0:
                set_cell_shading(cell, "F2F6F3")


def page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)


def build():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # =====================================================================
    # PAGE 1 - Product + value story
    # =====================================================================
    add_heading(doc, "FPO Integrated Operating System", size=22)
    add_para(
        doc,
        "An agent-led digital operating platform for Farmer Producer Organisations.",
        size=11, color=MUTED, italic=True, space_after=8,
    )

    add_subheading(doc, "What the Platform Is")
    add_para(
        doc,
        "Most agritech products solve one slice of the value chain - advisory, marketplace, "
        "satellite, or carbon. FPOs sit at the centre of the value chain and have to stitch "
        "these slices together manually across spreadsheets, phone calls, and WhatsApp groups. "
        "The FPO Integrated Operating System replaces that fragmentation with a single agent-led "
        "platform that runs the FPO's member registry, input fulfillment, produce collection, "
        "market execution, communication, and carbon readiness from one workflow spine.",
    )
    add_para(
        doc,
        "The system is built around a simple operating principle: autonomous by default, human "
        "only by exception. Specialist agents handle farmer messages, input demand, procurement, "
        "crop-cycle follow-up, harvest verification, and market allocation end to end whenever "
        "confidence is high and business rules permit. The FPO office is pulled in only for "
        "low-confidence signals, disputes, disease escalations, or transactions that cross "
        "approval thresholds.",
    )

    add_subheading(doc, "What FPOs Achieve With It")
    add_bullet(doc, "one source of truth for members, plots, crops, inventory, sales, and settlements.", bold_lead="Single operating layer: ")
    add_bullet(doc, "a farmer WhatsApp message for urea, harvest readiness, or price enquiry triggers demand capture, stock check, procurement, collection, and market allocation without manual re-entry.", bold_lead="Message-to-execution: ")
    add_bullet(doc, "the system watches crop seasons and nudges farmers as harvest windows approach, so collection and buyer allocation start on time.", bold_lead="Proactive outreach: ")
    add_bullet(doc, "procurement, sales allocation, and settlements flow to approvers only when thresholds are crossed, with a full audit trail of agent and human decisions.", bold_lead="Governed autonomy: ")
    add_bullet(doc, "live queues for fulfillment, harvest watchlist, market allocation, and human handoffs, so FPO managers see the autonomous operating state at a glance.", bold_lead="Command visibility: ")
    add_bullet(doc, "registry, plot, and practice discipline feeds directly into a carbon readiness layer, so climate revenue becomes feasible on the same data spine.", bold_lead="Carbon on the same data: ")

    add_subheading(doc, "Commercial Outcomes")
    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = "Light Grid"
    data = [
        ("Outcome", "How the platform delivers it"),
        ("Lower coordination cost", "Agents replace manual triaging of farmer messages, demand aggregation, and buyer matching. FPO staff move from data-entry work to exception handling."),
        ("Better procurement economics", "Compatible demands are grouped into autonomous purchase requests against the best-fit supplier; high-value POs route to approvers instead of bypassing governance."),
        ("Better sale realisation", "Produce collections are auto-matched against open buyer demand and mandi snapshots, so the FPO acts on the best current offer rather than the first offer."),
        ("Carbon readiness", "Practice records, plot boundaries, and crop seasons are captured as first-class entities, making the FPO eligible for aggregation into carbon projects."),
    ]
    for i, (a, b) in enumerate(data):
        tbl.rows[i].cells[0].text = a
        tbl.rows[i].cells[1].text = b
    style_table(tbl)

    add_subheading(doc, "Who It Is For")
    add_para(
        doc,
        "Built for FPO management, field coordinators, procurement and warehouse staff, and "
        "finance users. Farmers interact through a WhatsApp-style channel; buyers and suppliers "
        "plug in through the same operational records. The platform is the digital layer the FPO "
        "runs on, not another farmer-facing app.",
    )

    page_break(doc)

    # =====================================================================
    # PAGE 2 - System walkthrough (agentic)
    # =====================================================================
    add_heading(doc, "System Walkthrough", size=20)
    add_para(
        doc,
        "How farmer signals become business execution through a team of specialist agents.",
        size=11, color=MUTED, italic=True, space_after=8,
    )

    add_subheading(doc, "Specialist Agent Roster")
    tbl = doc.add_table(rows=6, cols=2)
    tbl.style = "Light Grid"
    rows = [
        ("Agent", "Responsibility"),
        ("Farmer Intake Agent", "Classifies every inbound farmer message into price, input, disease, harvest, advisory, or acknowledgement intents and scores confidence."),
        ("Input Fulfillment Agent", "Handles input demand end to end: stock check, supplier selection, autonomous purchase request, goods receipt, and farmer stock issue."),
        ("Crop Cycle Agent", "Watches sowing and expected harvest dates, opens proactive harvest-readiness conversations, and keeps collection recency in check."),
        ("Market Allocation Agent", "Pairs collections with the best open buyer demand, creates sales orders, and raises dispatches within the allocation policy."),
        ("Human Exception Agent", "Centralises handoffs: low-confidence cases, disputes, disease escalations, and anything crossing an approval threshold."),
    ]
    for i, (a, b) in enumerate(rows):
        tbl.rows[i].cells[0].text = a
        tbl.rows[i].cells[1].text = b
    style_table(tbl)

    add_subheading(doc, "End-to-End Runtime Flows")

    add_para(doc, "1. Input request (Farmer Intake + Input Fulfillment)", bold=True, size=10, color=ACCENT)
    add_bullet(doc, "Farmer: \"Need 3 bags urea for my crop.\" Intent classifier tags input_request and scores item, quantity, and crop context.")
    add_bullet(doc, "A structured input-demand record is created. Low confidence is marked needs_review and routed to the Human Exception Agent.")
    add_bullet(doc, "If stock is present, the agent issues the stock to the farmer directly and updates inventory and demand status.")
    add_bullet(doc, "If stock is absent, compatible demands are grouped into an autonomous purchase request against the best-fit supplier.")
    add_bullet(doc, "If the PR crosses the configured quantity or value thresholds, it is routed to Approvals; otherwise it proceeds and a goods receipt is raised on arrival.")
    add_bullet(doc, "Farmer receives a reply with the resulting state: stock issued, procurement started, or pending review.")

    add_para(doc, "2. Harvest readiness (Farmer Intake + Crop Cycle + Market Allocation)", bold=True, size=10, color=ACCENT)
    add_bullet(doc, "Farmer: \"My crop is harvest ready now.\" Intent classifier tags harvest_update.")
    add_bullet(doc, "Crop Cycle Agent scores the signal using keyword strength, crop season timing, and expected harvest window to produce a harvest_signal record.")
    add_bullet(doc, "On high confidence a produce collection is created against the farmer's active season at an estimated quantity.")
    add_bullet(doc, "Market Allocation Agent finds the best open buyer demand for that crop and creates a sales order against the collection.")
    add_bullet(doc, "If the sales order is within policy, a dispatch is raised immediately; if not, the order is held for approval.")
    add_bullet(doc, "Farmer gets a single reply covering verification, allocation, and dispatch status.")

    add_para(doc, "3. Proactive crop-cycle outreach (Crop Cycle)", bold=True, size=10, color=ACCENT)
    add_bullet(doc, "Orchestration cycle scans active crop seasons against a harvest outreach window and an alert cooldown.")
    add_bullet(doc, "Seasons approaching harvest with no recent collection and no recent alert trigger a proactive harvest-check message.")
    add_bullet(doc, "Each outreach is logged as an agent alert and task, so telemetry shows the system is reaching out, not just reacting.")

    add_para(doc, "4. Price enquiry and advisory (Farmer Intake)", bold=True, size=10, color=ACCENT)
    add_bullet(doc, "Price queries are answered directly from the latest crop-mandi snapshot; missing data triggers a handoff instead of a fabricated reply.")
    add_bullet(doc, "Advisory queries use crop reference data and active season timing to respond with concise stage-appropriate guidance.")

    add_subheading(doc, "Command Centre and Governance")
    add_para(
        doc,
        "The Command screen is the operating surface for the autonomous system: agent performance summary, "
        "fulfillment queue, harvest watchlist, market queue, human handoffs, recent runs, alerts, and tasks. "
        "The Approvals screen is where governance lives - purchase requests, sales orders, settlements, and "
        "any agent-initiated action that crosses a threshold are released or blocked here. The Handoff Desk "
        "collects low-confidence or escalated conversations so staff attention is concentrated on judgement work.",
    )

    add_subheading(doc, "Policy Thresholds That Keep Autonomy Safe")
    add_bullet(doc, "trust floor below which input demand must be reviewed.", bold_lead="Input autonomy trust: ")
    add_bullet(doc, "score floor below which harvest messages cannot trigger collection.", bold_lead="Harvest confidence: ")
    add_bullet(doc, "quantity and value ceilings beyond which a PR must be approved.", bold_lead="PR auto-approval: ")
    add_bullet(doc, "quantity ceiling beyond which a sales order holds for approval.", bold_lead="Sales order approval: ")
    add_bullet(doc, "how close to expected harvest outreach fires, and how long before the same farmer can be nudged again.", bold_lead="Outreach window and cooldown: ")

    page_break(doc)

    # =====================================================================
    # PAGE 3 - Technical manual
    # =====================================================================
    add_heading(doc, "Technical Manual", size=20)
    add_para(
        doc,
        "Architecture, data requirements, and technology stack.",
        size=11, color=MUTED, italic=True, space_after=8,
    )

    add_subheading(doc, "Architecture at a Glance")
    add_bullet(doc, "React single-page application rendering the Command centre, Farmer Network, Fulfillment, Market Execution, Handoff Desk, Campaigns, Approvals, and Carbon modules.", bold_lead="Frontend: ")
    add_bullet(doc, "Flask API service exposing registry, operations, market, communication, carbon, approvals, and agent-orchestration endpoints over a consistent JSON contract.", bold_lead="Backend: ")
    add_bullet(doc, "A dataset module holds the operational records and a rule-based orchestration layer reads and writes into the same records used by every module, keeping every screen in sync.", bold_lead="Data and orchestration: ")
    add_bullet(doc, "Each orchestration pass is recorded under agent_runs, agent_tasks, agent_alerts, and harvest_signals tables, and every business record carries source, source_ref, and created_by_agent provenance for audit.", bold_lead="Observability: ")

    add_subheading(doc, "Technology Stack")
    tbl = doc.add_table(rows=6, cols=2)
    tbl.style = "Light Grid"
    rows = [
        ("Layer", "Technology"),
        ("Frontend", "React 18, Vite 5, modular section views with shared StatCard, DataTable, and SectionPanel primitives."),
        ("Backend", "Python, Flask 3, Flask-CORS, modular route groups for registry, operations, market, communication, carbon, admin, and agent orchestration."),
        ("Data layer", "In-memory dataset generated deterministically by seed, with cross-entity reconciliation rules so every transaction links back to a prior state."),
        ("Agent layer", "Rule-based orchestration with explicit thresholds, intent classification, confidence scoring, and autonomous helpers for procurement, collection, sales, and dispatch."),
        ("Integration surfaces", "Mock WhatsApp send and reply endpoints, broadcast simulation, agent command-centre API, and cycle trigger API ready to be wired to real channels."),
    ]
    for i, (a, b) in enumerate(rows):
        tbl.rows[i].cells[0].text = a
        tbl.rows[i].cells[1].text = b
    style_table(tbl)

    add_subheading(doc, "Core Data Domains")
    tbl = doc.add_table(rows=11, cols=2)
    tbl.style = "Light Grid"
    rows = [
        ("Domain", "Key fields"),
        ("Geography", "state, district, block, village hierarchy"),
        ("FPO master", "FPO name, crop focus, member count, warehouse capacity"),
        ("Farmer master", "farmer id, village, contact, language, land size, irrigation, soil"),
        ("Plot and season", "plot id, area, boundary, soil, irrigation; crop, sowing date, expected harvest, variety"),
        ("Communication", "phone, language, preferred channel, message logs, advisory logs, escalations, broadcasts"),
        ("Input and procurement", "input catalog, supplier master, input demand, purchase request, purchase order, goods receipt, inventory, input issue"),
        ("Produce and settlement", "produce collection, grade, quantity, centre, farmer settlement"),
        ("Market", "buyer master, buyer demand, mandi price time series, sales order, dispatch"),
        ("Carbon", "practice records, carbon factors, estimates, aggregated carbon projects"),
        ("Agent traces", "agent_runs, agent_tasks, agent_alerts, harvest_signals, plus source provenance on every business record"),
    ]
    for i, (a, b) in enumerate(rows):
        tbl.rows[i].cells[0].text = a
        tbl.rows[i].cells[1].text = b
    style_table(tbl)

    add_subheading(doc, "Primary API Surface")
    add_bullet(doc, "/api/agent/command-center, /api/agent/run-cycle, /api/agent/runs/<id>/events", bold_lead="Agent orchestration: ")
    add_bullet(doc, "/api/registry/fpos, /farmers, /plots, /seasons, /communication-profiles, /geographies", bold_lead="Registry: ")
    add_bullet(doc, "/api/operations/demand-summary, /demands/review-queue, /procurement, /purchase-requests, /goods-receipts, /inventory, /input-issues, /collections, /settlements", bold_lead="Operations: ")
    add_bullet(doc, "/api/market/prices, /buyers, /demands, /matching, /sales-orders, /dispatches", bold_lead="Market: ")
    add_bullet(doc, "/api/communication/inbox, /advisories, /escalations, /disease-cases, /broadcasts, /mock-whatsapp/send, /mock-whatsapp/reply, /agent-reply", bold_lead="Communication: ")
    add_bullet(doc, "/api/carbon/practices, /estimates, /projects", bold_lead="Carbon: ")
    add_bullet(doc, "/api/admin/roles, /audit-logs, /approval-logs, /approvals/<id>/decide, /approvals/bulk-decide", bold_lead="Governance: ")

    add_subheading(doc, "Data Requirements for Deployment")
    add_para(doc, "Priority 1 - operationally critical", bold=True, size=10, color=ACCENT)
    add_bullet(doc, "FPO profile and governance structure, member registry, village mapping.")
    add_bullet(doc, "Crop and season records per plot, input catalog and supplier master.")
    add_bullet(doc, "Produce collection formats, settlement process, approval policy.")
    add_para(doc, "Priority 2 - market linkage", bold=True, size=10, color=ACCENT)
    add_bullet(doc, "Mandi and buyer price feeds, buyer and processor directory.")
    add_bullet(doc, "Historical FPO sales, payment reliability, logistics constraints.")
    add_para(doc, "Priority 3 - climate and land intelligence", bold=True, size=10, color=ACCENT)
    add_bullet(doc, "Plot boundaries, soil and irrigation attributes, climate practice records.")
    add_bullet(doc, "Satellite-ready identifiers, approved carbon methodology assumptions.")

    add_subheading(doc, "Security, Audit, and Extensibility")
    add_bullet(doc, "Role-based views, approval queues, and audit logs are first-class - every autonomous and human decision is traced.")
    add_bullet(doc, "Provenance fields (source, source_ref, created_by_agent) allow operators to distinguish office-created, agent-created, and message-originated records.")
    add_bullet(doc, "The orchestration cycle is explicit and triggerable, so it can be moved to a scheduled worker without changing the business logic.")
    add_bullet(doc, "All external signals - WhatsApp, mandi feeds, satellite, carbon factors - are behind stable interfaces that can be swapped for live integrations without touching module code.")

    doc.save(r"c:\\Users\\AkashPatil\\POC_FPO\\FPO_Integrated_OS_Brief.docx")
    print("OK")


if __name__ == "__main__":
    build()
